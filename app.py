import streamlit as st
import urllib3
import json
import io

# URL API Groq (Jaringan Stabil & Super Cepat)
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"

# --- KONFIGURASI HALAMAN BROWSER ---
st.set_page_config(
    page_title="Generator TP & ATP Resmi", 
    page_icon="🌋", 
    layout="wide"
)

# --- SENTUHAN CSS KUSTOM (GRADASI TEMA ALAM GUNUNG DUKONO - ANTI BLANK) ---
st.markdown("""
    <style>
        /* Gaya font dan warna latar belakang dashboard utama */
        .main { background-color: #f8fafc; }
        
        /* Banner Atas: Gradasi Estetik Abu-abu Vulkanik & Hijau Hutan Halmahera Utara */
        .welcome-card {
            background: linear-gradient(135deg, #2c3e50 0%, #1e293b 40%, #115e59 100%);
            color: white;
            padding: 60px 40px;
            border-radius: 16px;
            box-shadow: 0 10px 25px -5px rgba(17, 94, 89, 0.3);
            margin-bottom: 30px;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            text-align: center;
        }
        
        /* Kartu Konten Input Elemen */
        .element-box {
            background-color: white;
            padding: 25px;
            border-radius: 12px;
            border-left: 5px solid #10b981;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
            margin-bottom: 20px;
        }
    </style>
""", unsafe_allow_html=True)

# --- AREA HEADER UTAMA RATA TENGAH SEMPURNA (BERSIH & MINIMALIS) ---
st.markdown("""
    <div class="welcome-card">
        <h1 style='margin: 0 auto 15px auto; font-weight: 800; color: white; text-shadow: 2px 2px 4px rgba(0,0,0,0.5); font-size: 36px; letter-spacing: 0.5px; text-align: center; width: 100%;'>
            Generator TP & ATP Kurikulum Merdeka
        </h1>
        <p style='font-size: 16px; font-weight: 400; color: #ccfbf1; margin: 0 auto; text-shadow: 1px 1px 2px rgba(0,0,0,0.5); max-width: 800px; text-align: center; line-height: 1.6;'>
            Asisten cerdas perancang administrasi sekolah terpadu.
        </p>
    </div>
""", unsafe_allow_html=True)

# --- FORM INPUT UTAMA ---

# Section 1: Validasi API Key
st.markdown("### 🔑 Langkah 1: Validasi API Key")
api_key_input = st.text_input("Masukkan Groq API Key Anda (diawali gsk_...):", type="password", placeholder="gsk_...")
st.caption("Dapatkan kunci gratis Anda melalui situs resmi [console.groq.com](https://console.groq.com)")

st.markdown("<br>", unsafe_allow_html=True)

# Section 2: Informasi Umum Kurikulum
st.markdown("### 🏫 Langkah 2: Informasi Umum")
col_meta1, col_meta2 = st.columns(2)
with col_meta1:
    mata_pelajaran = st.text_input("Mata Pelajaran:", placeholder="Contoh: Matematika")
with col_meta2:
    fase_kelas = st.text_input("Fase / Kelas:", placeholder="Contoh: A / Kelas 1")

st.markdown("<br>", unsafe_allow_html=True)

# Section 3: Manajemen Multi-Elemen (Session State)
st.markdown("### 📋 Langkah 3: Input Elemen & Capaian Pembelajaran (CP)")
st.info("💡 Anda bisa memasukkan lebih dari satu elemen sekaligus agar otomatis digabungkan ke dalam satu tabel dokumen resmi.")

if "list_elemen" not in st.session_state:
    st.session_state.list_elemen = [{"elemen": "", "cp": ""}]

# Looping render kotak elemen di halaman utama
for index, item in enumerate(st.session_state.list_elemen):
    st.markdown(f'<div class="element-box">', unsafe_allow_html=True)
    
    col_el, col_btn = st.columns([6, 1])
    with col_el:
        st.markdown(f"**🔸 Form Pengisian Elemen ke-{index + 1}**")
        el_name = st.text_input(f"Nama Elemen ke-{index + 1}:", value=item['elemen'], key=f"el_{index}", placeholder="Contoh: Bilangan / Aljabar / Geometri")
        cp_text = st.text_area(f"Kalimat Capaian Pembelajaran (CP) ke-{index + 1}:", value=item['cp'], key=f"cp_{index}", height=100, placeholder="Tempel kalimat CP asli di sini...")
        
        # Simpan data secara realtime ke memori aplikasi
        st.session_state.list_elemen[index]['elemen'] = el_name
        st.session_state.list_elemen[index]['cp'] = cp_text
    with col_btn:
        st.write("<br>", unsafe_allow_html=True)
        if st.button("🗑️ Hapus", key=f"del_{index}", use_container_width=True):
            if len(st.session_state.list_elemen) > 1:
                st.session_state.list_elemen.pop(index)
                st.rerun()
            else:
                st.warning("Minimal harus ada 1 elemen input!")
                
    st.markdown('</div>', unsafe_allow_html=True)

# Tombol Tambah Elemen Baru
col_tambah, _ = st.columns([2, 5])
with col_tambah:
    if st.button("➕ Tambah Elemen & Kalimat CP", use_container_width=True):
        st.session_state.list_elemen.append({"elemen": "", "cp": ""})
        st.rerun()

st.markdown("<br><hr><br>", unsafe_allow_html=True)

# --- FUNGSI EXPORT MICROSOFT WORD ---
def buat_dokumen_word(data_gabungan, mata_pelajaran, fase_kelas):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    
    p_title = doc.add_paragraph()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("DOKUMEN RANCANGAN MATERI, TP, DAN ALUR TUJUAN PEMBELAJARAN (ATP)")
    run_title.bold = True
    run_title.size = Pt(14)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"Mata Pelajaran: {mata_pelajaran} | Fase: {fase_kelas}")
    run_sub.size = Pt(11)
    doc.add_paragraph("")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ["Elemen", "Capaian Pembelajaran (CP)", "Tujuan Pembelajaran (TP)", "Alur Tujuan Pembelajaran (ATP)"]
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc_pr = hdr_cells[idx]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), "4F46E5")
        tc_pr.append(shd)
        for p in hdr_cells[idx].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                
    for data_json in data_gabungan:
        row_cells = table.add_row().cells
        row_cells[0].text = data_json.get("elemen", "")
        row_cells[1].text = data_json.get("capaian_pembelajaran", "")
        
        tp_lines = [f"[{tp.get('tp_id', '')}] {tp.get('indikator_kompetensi', '')}" for tp in data_json.get("tujuan_pembelajaran_list", [])]
        row_cells[2].text = "\n".join(tp_lines)
        
        atp_lines = []
        for alur in data_json.get("alur_tujuan_pembelajaran", []):
            for step_item in alur.get("urutan_tp", []):
                atp_lines.append(step_item.get("kalimat_alur", ""))
        row_cells[3].text = "\n".join(atp_lines)
        
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- FUNGSI EXPORT PDF RESMI ---
def buat_dokumen_pdf(data_gabungan, mata_pelajaran, fase_kelas):
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], alignment=1, fontSize=14, spaceAfter=6)
    sub_style = ParagraphStyle('SubStyle', parent=styles['Normal'], alignment=1, fontSize=11, spaceAfter=15)
    cell_style = ParagraphStyle('CellStyle', parent=styles['Normal'], fontSize=9, leading=12)
    header_style = ParagraphStyle('HeaderStyle', parent=styles['Normal'], fontSize=10, fontName="Helvetica-Bold", textColor=colors.white)
    
    story.append(Paragraph("<b>DOKUMEN RANCANGAN MATERI, TP, DAN ALUR TUJUAN PEMBELAJARAN (ATP)</b>", title_style))
    story.append(Paragraph(f"Mata Pelajaran: {mata_pelajaran} | Fase: {fase_kelas}", sub_style))
    
    table_data = [[
        Paragraph("<b>Elemen</b>", header_style),
        Paragraph("<b>Capaian Pembelajaran (CP)</b>", header_style),
        Paragraph("<b>Tujuan Pembelajaran (TP)</b>", header_style),
        Paragraph("<b>Alur Tujuan Pembelajaran (ATP)</b>", header_style)
    ]]
    
    for data_json in data_gabungan:
        tp_lines = [f"[{tp.get('tp_id', '')}] {tp.get('indikator_kompetensi', '')}" for tp in data_json.get("tujuan_pembelajaran_list", [])]
        atp_lines = []
        for alur in data_json.get("alur_tujuan_pembelajaran", []):
            for step_item in alur.get("urutan_tp", []):
                atp_lines.append(step_item.get("kalimat_alur", ""))
                
        table_data.append([
            Paragraph(data_json.get("elemen", ""), cell_style),
            Paragraph(data_json.get("capaian_pembelajaran", "").replace("\n", "<br/>"), cell_style),
            Paragraph("<br/>".join(tp_lines), cell_style),
            Paragraph("<br/>".join(atp_lines), cell_style)
        ])
    
    col_widths = [100, 190, 220, 220]
    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4F46E5')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    
    story.append(t)
    doc.build(story)
    buffer.seek(0)
    return buffer

# --- PROSES INTEGRASI AI ---
col_aksi, _ = st.columns([3, 4])
with col_aksi:
    tombol_proses = st.button("🚀 Buat CP & ATP", type="primary", use_container_width=True)

if tombol_proses:
    api_key = api_key_input.strip()
    ada_data_kosong = any([not item['elemen'] or not item['cp'] for item in st.session_state.list_elemen])
    
    if not api_key:
        st.error("Silakan masukkan Groq API Key terlebih dahulu!")
    elif not mata_pelajaran or not fase_kelas or ada_data_kosong:
        st.error("Mohon lengkapi semua data Informasi Umum beserta seluruh isi kolom Elemen & CP!")
    else:
        with st.spinner("⏳ AI sedang merumuskan kode desimal kurikulum Anda..."):
            
            singkatan_mapel = "".join([kata[0] for kata in mata_pelajaran.split()]).upper()
            if len(singkatan_mapel) < 3:
                singkatan_mapel = mata_pelajaran[:3].upper()
            fase_clean = fase_kelas.split("/")[0].strip().upper()[:1]
            
            data_hasil_gabungan = []
            proses_sukses = True
            
            for idx, item in enumerate(st.session_state.list_elemen):
                nomor_elemen = idx + 1
                
                prompt = f"""Anda adalah pakar kurikulum Kurikulum Merdeka. Analisis data elemen berikut:
                Mata Pelajaran: {mata_pelajaran}
                Fase/Kelas: {fase_kelas}
                Elemen ke-{nomor_elemen}: {item['elemen']}
                Capaian Pembelajaran: {item['cp']}
                
                Tugas Utama Anda:
                1. Bedah kalimat CP tersebut untuk memisahkan KOMPETENSI dan KONTEN INTI.
                2. Rumuskan daftar Tujuan Pembelajaran (TP) pada array 'tujuan_pembelajaran_list'.
                   ATURAN KODE WAJIB MENGGUNAKAN TANDA TITIK DESIMAL (Elemen.Urutan).
                3. SUSUN ATP SECARA HIERARKIS: Urutkan langkah lintas konsep secara runtut dari prasyarat dasar ke kompleks.
                4. Pada kunci 'kalimat_alur', gabungkan nomor urut langkah, ID TP desimal dinamis, dan kalimat kompetensinya.

                Hanya keluarkan hasil dalam format JSON objek mentah yang valid tanpa pembungkus teks markdown. Struktur mutlak:
                {{
                  "elemen": "{item['elemen']}",
                  "capaian_pembelajaran": "{item['cp']}",
                  "tujuan_pembelajaran_list": [
                    {{
                      "tp_id": "TP-{singkatan_mapel}-{fase_clean}-{nomor_elemen}.1",
                      "indikator_kompetensi": "Peserta didik mampu..."
                    }}
                  ],
                  "alur_tujuan_pembelajaran": [
                    {{
                      "atp_id": "ATP-{singkatan_mapel}-{fase_clean}-{nomor_elemen}-2026",
                      "nama_alur": "Alur Fase {fase_kelas}",
                      "urutan_tp": [
                        {{
                          "step": 1,
                          "tp_id": "TP-{singkatan_mapel}-{fase_clean}-{nomor_elemen}.1",
                          "kalimat_alur": "1. [TP-{singkatan_mapel}-{fase_clean}-{nomor_elemen}.1] Peserta didik mampu..."
                        }}
                      ]
                    }}
                  ]
                }}"""
                
                try:
                    http = urllib3.PoolManager()
                    payload = json.dumps({
                        "model": "llama-3.3-70b-versatile",
                        "messages": [
                            {"role": "system", "content": "You are a professional curriculum builder. You must respond ONLY with a raw, valid JSON object."},
                            {"role": "user", "content": prompt}
                        ],
                        "temperature": 0.1,
                        "response_format": {"type": "json_object"}
                    })
                    
                    response = http.request('POST', GROQ_URL, body=payload, headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}, timeout=30.0)
                    
                    if response.status == 200:
                        res_json = json.loads(response.data.decode('utf-8'))
                        raw_text = res_json['choices'][0]['message']['content'].strip()
                        data_json = json.loads(raw_text)
                        data_hasil_gabungan.append(data_json)
                    else:
                        proses_sukses = False
                        st.error(f"Gagal memproses elemen '{item['elemen']}'. Status Server: {response.status}")
                        break
                except Exception as e:
                    proses_sukses = False
                    st.error(f"Terjadi kesalahan di elemen '{item['elemen']}': {str(e)}")
                    break
            
            if proses_sukses and data_hasil_gabungan:
                st.balloons()
                st.success(f"🎉 Hebat! Berhasil memproses total {len(data_hasil_gabungan)} elemen!")
                
                st.markdown("### 📥 Ambil Dokumen Hasil Unduhan Resmi")
                col_dl1, col_dl2 = st.columns(2)
                
                with col_dl1:
                    word_file = buat_dokumen_word(data_hasil_gabungan, mata_pelajaran, fase_kelas)
                    st.download_button(
                        label="📄 Download File Microsoft Word (.docx)",
                        data=word_file,
                        file_name=f"Rancangan_Kurikulum_{singkatan_mapel}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
                with col_dl2:
                    pdf_file = buat_dokumen_pdf(data_hasil_gabungan, mata_pelajaran, fase_kelas)
                    st.download_button(
                        label="📕 Download File PDF Landscape (.pdf)",
                        data=pdf_file,
                        file_name=f"Rancangan_Kurikulum_{singkatan_mapel}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                
                st.markdown("---")
                st.subheader("📊 Pratinjau Struktur JSON")
                st.json(data_hasil_gabungan)